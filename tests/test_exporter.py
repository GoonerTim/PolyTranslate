"""Tests for TranslationExporter."""

from __future__ import annotations

import xml.etree.ElementTree as ET
from pathlib import Path

import pytest

from app.core.exporter import TranslationExporter

NS = "urn:oasis:names:tc:xliff:document:1.2"
_ns = f"{{{NS}}}"


@pytest.fixture
def sample_data() -> dict:
    return {
        "original_text": "Hello world.\nThis is a test.",
        "translations": {
            "deepl": "Hallo Welt.\nDies ist ein Test.",
            "google": "Hallo Welt.\nDas ist ein Test.",
        },
        "source_lang": "en",
        "target_lang": "de",
        "file_name": "test.txt",
    }


class TestExporterDocx:
    def test_export_docx(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.docx"
        result = TranslationExporter.export(
            **sample_data,
            output_path=output,
        )
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_export_docx_content(self, tmp_path: Path, sample_data: dict) -> None:
        from docx import Document

        output = tmp_path / "result.docx"
        TranslationExporter.export(**sample_data, output_path=output)

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "PolyTranslate Export" in full_text
        assert "EN" in full_text
        assert "DE" in full_text
        assert "Hello world." in full_text
        assert "Hallo Welt." in full_text
        assert "DEEPL" in full_text
        assert "GOOGLE" in full_text

    def test_export_docx_with_file_name(self, tmp_path: Path, sample_data: dict) -> None:
        from docx import Document

        output = tmp_path / "result.docx"
        TranslationExporter.export(**sample_data, output_path=output)

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "test.txt" in full_text

    def test_export_docx_no_file_name(self, tmp_path: Path, sample_data: dict) -> None:
        sample_data["file_name"] = ""
        output = tmp_path / "result.docx"
        result = TranslationExporter.export(**sample_data, output_path=output)
        assert result.exists()

    def test_export_docx_creates_parent_dirs(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "sub" / "dir" / "result.docx"
        result = TranslationExporter.export(**sample_data, output_path=output)
        assert result.exists()

    def test_export_docx_single_service(self, tmp_path: Path) -> None:
        output = tmp_path / "result.docx"
        result = TranslationExporter.export(
            original_text="Hello",
            translations={"google": "Hola"},
            source_lang="en",
            target_lang="es",
            output_path=output,
        )
        assert result.exists()


class TestExporterPdf:
    def test_export_pdf(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.pdf"
        result = TranslationExporter.export(**sample_data, output_path=output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_export_pdf_is_valid(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.pdf"
        TranslationExporter.export(**sample_data, output_path=output)
        content = output.read_bytes()
        assert content.startswith(b"%PDF")

    def test_export_pdf_no_file_name(self, tmp_path: Path, sample_data: dict) -> None:
        sample_data["file_name"] = ""
        output = tmp_path / "result.pdf"
        result = TranslationExporter.export(**sample_data, output_path=output)
        assert result.exists()

    def test_export_pdf_creates_parent_dirs(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "nested" / "result.pdf"
        result = TranslationExporter.export(**sample_data, output_path=output)
        assert result.exists()

    def test_export_pdf_special_chars(self, tmp_path: Path) -> None:
        output = tmp_path / "result.pdf"
        result = TranslationExporter.export(
            original_text='Test <html> & "quotes"',
            translations={"deepl": 'Test <html> & "Anf\u00fchrungszeichen"'},
            source_lang="en",
            target_lang="de",
            output_path=output,
        )
        assert result.exists()

    def test_export_pdf_empty_lines(self, tmp_path: Path) -> None:
        output = tmp_path / "result.pdf"
        result = TranslationExporter.export(
            original_text="Line 1\n\nLine 3",
            translations={"google": "Zeile 1\n\nZeile 3"},
            source_lang="en",
            target_lang="de",
            output_path=output,
        )
        assert result.exists()


class TestExporterXliff:
    def test_export_xliff(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.xliff"
        result = TranslationExporter.export(**sample_data, output_path=output)
        assert result == output
        assert output.exists()

    def test_export_xliff_structure(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.xliff"
        TranslationExporter.export(**sample_data, output_path=output)

        tree = ET.parse(str(output))
        root = tree.getroot()

        assert root.tag.endswith("xliff")
        assert root.attrib["version"] == "1.2"

        files = root.findall(f"{_ns}file")
        assert len(files) == 2  # deepl + google

        for file_elem in files:
            assert file_elem.attrib["source-language"] == "en"
            assert file_elem.attrib["target-language"] == "de"
            assert file_elem.attrib["original"] == "test.txt"

    def test_export_xliff_trans_units(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.xliff"
        TranslationExporter.export(**sample_data, output_path=output)

        tree = ET.parse(str(output))
        root = tree.getroot()

        first_file = root.findall(f"{_ns}file")[0]
        body = first_file.find(f"{_ns}body")
        assert body is not None
        units = body.findall(f"{_ns}trans-unit")

        assert len(units) == 2
        assert units[0].find(f"{_ns}source").text == "Hello world."  # type: ignore[union-attr]
        assert units[0].find(f"{_ns}target").text == "Hallo Welt."  # type: ignore[union-attr]
        assert units[1].find(f"{_ns}source").text == "This is a test."  # type: ignore[union-attr]
        assert units[1].find(f"{_ns}target").text == "Dies ist ein Test."  # type: ignore[union-attr]

    def test_export_xliff_service_tool_id(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.xliff"
        TranslationExporter.export(**sample_data, output_path=output)

        tree = ET.parse(str(output))
        root = tree.getroot()
        files = root.findall(f"{_ns}file")

        tool_ids = [f.attrib["tool-id"] for f in files]
        assert "polytranslate-deepl" in tool_ids
        assert "polytranslate-google" in tool_ids

    def test_export_xlf_extension(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.xlf"
        result = TranslationExporter.export(**sample_data, output_path=output)
        assert result.exists()

        tree = ET.parse(str(output))
        root = tree.getroot()
        assert root.tag.endswith("xliff")

    def test_export_xliff_skips_empty_lines(self, tmp_path: Path) -> None:
        output = tmp_path / "result.xliff"
        TranslationExporter.export(
            original_text="Line 1\n\nLine 3",
            translations={"google": "Zeile 1\n\nZeile 3"},
            source_lang="en",
            target_lang="de",
            output_path=output,
        )

        tree = ET.parse(str(output))
        root = tree.getroot()
        body = root.findall(f"{_ns}file")[0].find(f"{_ns}body")
        assert body is not None
        units = body.findall(f"{_ns}trans-unit")
        assert len(units) == 2  # empty line skipped

    def test_export_xliff_no_file_name(self, tmp_path: Path, sample_data: dict) -> None:
        sample_data["file_name"] = ""
        output = tmp_path / "result.xliff"
        TranslationExporter.export(**sample_data, output_path=output)

        tree = ET.parse(str(output))
        root = tree.getroot()
        file_elem = root.findall(f"{_ns}file")[0]
        assert file_elem.attrib["original"] == "translation"


class TestExporterCommon:
    def test_unsupported_format(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.xyz"
        with pytest.raises(ValueError, match="Unsupported export format"):
            TranslationExporter.export(**sample_data, output_path=output)

    def test_explicit_format_override(self, tmp_path: Path, sample_data: dict) -> None:
        output = tmp_path / "result.txt"
        result = TranslationExporter.export(**sample_data, output_path=output, fmt="docx")
        assert result.exists()

    def test_format_detection_from_extension(self, tmp_path: Path, sample_data: dict) -> None:
        for ext in ("docx", "pdf", "xliff", "xlf"):
            output = tmp_path / f"result.{ext}"
            result = TranslationExporter.export(**sample_data, output_path=output)
            assert result.exists()

    def test_export_with_path_string(self, tmp_path: Path, sample_data: dict) -> None:
        output = str(tmp_path / "result.docx")
        result = TranslationExporter.export(**sample_data, output_path=output)
        assert result.exists()
