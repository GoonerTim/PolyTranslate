"""Export translations to DOCX, PDF, and XLIFF formats."""

from __future__ import annotations

import logging
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class TranslationExporter:
    """Exports original + translated text to DOCX, PDF, and XLIFF formats."""

    SUPPORTED_FORMATS = {"docx", "pdf", "xliff"}

    @staticmethod
    def export(
        original_text: str,
        translations: dict[str, str],
        source_lang: str,
        target_lang: str,
        output_path: str | Path,
        fmt: str | None = None,
        file_name: str = "",
    ) -> Path:
        output_path = Path(output_path)
        if fmt is None:
            suffix = output_path.suffix.lstrip(".")
            fmt = "xliff" if suffix in ("xlf", "xliff") else suffix

        if fmt not in TranslationExporter.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported export format: {fmt}. "
                f"Supported: {', '.join(sorted(TranslationExporter.SUPPORTED_FORMATS))}"
            )

        if fmt == "docx":
            return TranslationExporter._export_docx(
                original_text, translations, source_lang, target_lang, output_path, file_name
            )
        elif fmt == "pdf":
            return TranslationExporter._export_pdf(
                original_text, translations, source_lang, target_lang, output_path, file_name
            )
        else:
            return TranslationExporter._export_xliff(
                original_text, translations, source_lang, target_lang, output_path, file_name
            )

    @staticmethod
    def _export_docx(
        original_text: str,
        translations: dict[str, str],
        source_lang: str,
        target_lang: str,
        output_path: Path,
        file_name: str,
    ) -> Path:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        doc = Document()

        # Title
        title = doc.add_heading("PolyTranslate Export", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(f"{source_lang.upper()} \u2192 {target_lang.upper()}")
        meta_run.font.size = Pt(12)
        meta_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        if file_name:
            meta_para.add_run(f"  |  {file_name}").font.size = Pt(12)

        doc.add_paragraph()

        # Original text section
        doc.add_heading("Original Text", level=1)
        for line in original_text.split("\n"):
            para = doc.add_paragraph(line)
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Translation sections
        for service, translation in translations.items():
            doc.add_page_break()
            heading = doc.add_heading(f"Translation: {service.upper()}", level=1)
            # Color the heading
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

            for line in translation.split("\n"):
                doc.add_paragraph(line)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("Exported DOCX to %s", output_path)
        return output_path

    @staticmethod
    def _export_pdf(
        original_text: str,
        translations: dict[str, str],
        source_lang: str,
        target_lang: str,
        output_path: Path,
        file_name: str,
    ) -> Path:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as err:
            raise ImportError(
                "reportlab is required for PDF export. Install it with: pip install reportlab"
            ) from err

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ExportTitle",
            parent=styles["Title"],
            fontSize=20,
            spaceAfter=6,
        )
        meta_style = ParagraphStyle(
            "ExportMeta",
            parent=styles["Normal"],
            fontSize=11,
            textColor="#606060",
            alignment=1,  # center
            spaceAfter=20,
        )
        heading_style = ParagraphStyle(
            "ExportHeading",
            parent=styles["Heading1"],
            fontSize=14,
            textColor="#2563eb",
            spaceBefore=16,
            spaceAfter=8,
        )
        original_heading_style = ParagraphStyle(
            "OriginalHeading",
            parent=styles["Heading1"],
            fontSize=14,
            textColor="#333333",
            spaceBefore=16,
            spaceAfter=8,
        )
        body_style = ParagraphStyle(
            "ExportBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=4,
        )

        story: list[Any] = []

        # Title
        story.append(Paragraph("PolyTranslate Export", title_style))
        meta_text = f"{source_lang.upper()} \u2192 {target_lang.upper()}"
        if file_name:
            meta_text += f"  |  {file_name}"
        story.append(Paragraph(meta_text, meta_style))

        # Original text
        story.append(Paragraph("Original Text", original_heading_style))
        for line in original_text.split("\n"):
            escaped = _escape_xml(line) if line.strip() else "&nbsp;"
            story.append(Paragraph(escaped, body_style))

        # Translations
        for service, translation in translations.items():
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"Translation: {service.upper()}", heading_style))
            for line in translation.split("\n"):
                escaped = _escape_xml(line) if line.strip() else "&nbsp;"
                story.append(Paragraph(escaped, body_style))

        doc.build(story)
        logger.info("Exported PDF to %s", output_path)
        return output_path

    @staticmethod
    def _export_xliff(
        original_text: str,
        translations: dict[str, str],
        source_lang: str,
        target_lang: str,
        output_path: Path,
        file_name: str,
    ) -> Path:
        nsmap = "urn:oasis:names:tc:xliff:document:1.2"
        root = ET.Element("xliff", version="1.2", xmlns=nsmap)

        original_lines = original_text.split("\n")

        for service, translation in translations.items():
            translated_lines = translation.split("\n")

            file_elem = ET.SubElement(
                root,
                "file",
                {
                    "original": file_name or "translation",
                    "source-language": source_lang,
                    "target-language": target_lang,
                    "datatype": "plaintext",
                    "tool-id": f"polytranslate-{service}",
                },
            )
            body = ET.SubElement(file_elem, "body")

            max_lines = max(len(original_lines), len(translated_lines))
            for i in range(max_lines):
                src_line = original_lines[i] if i < len(original_lines) else ""
                tgt_line = translated_lines[i] if i < len(translated_lines) else ""

                if not src_line.strip() and not tgt_line.strip():
                    continue

                trans_unit = ET.SubElement(body, "trans-unit", id=f"{service}-{i + 1}")
                source_elem = ET.SubElement(trans_unit, "source")
                source_elem.text = src_line
                target_elem = ET.SubElement(trans_unit, "target")
                target_elem.text = tgt_line

        output_path.parent.mkdir(parents=True, exist_ok=True)

        tree = ET.ElementTree(root)
        ET.indent(tree, space="  ")
        tree.write(str(output_path), encoding="unicode", xml_declaration=True)

        logger.info("Exported XLIFF to %s", output_path)
        return output_path


def _escape_xml(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
